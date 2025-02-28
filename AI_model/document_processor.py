import PyPDF2
import docx
import logging
from pathlib import Path
from typing import Union, Dict, Any
from model_config import DOC_CONFIG

class DocumentProcessor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.supported_formats = DOC_CONFIG['supported_formats']
        self.max_file_size = DOC_CONFIG['max_file_size']

    def process_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Process a document and extract its text content
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File {file_path} not found")
            
        if file_path.stat().st_size > self.max_file_size:
            raise ValueError(f"File size exceeds maximum allowed size of {self.max_file_size} bytes")
            
        if file_path.suffix not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        try:
            if file_path.suffix == '.pdf':
                return self._process_pdf(file_path)
            elif file_path.suffix in ['.doc', '.docx']:
                return self._process_word(file_path)
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            raise

    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF documents"""
        text_content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content += page.extract_text()
        
        return {
            'content': text_content,
            'metadata': {
                'pages': len(pdf_reader.pages),
                'format': 'pdf'
            }
        }

    def _process_word(self, file_path: Path) -> Dict[str, Any]:
        """Process Word documents"""
        doc = docx.Document(file_path)
        text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return {
            'content': text_content,
            'metadata': {
                'paragraphs': len(doc.paragraphs),
                'format': file_path.suffix[1:]
            }
        } 